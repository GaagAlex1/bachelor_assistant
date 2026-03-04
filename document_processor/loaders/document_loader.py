"""
DOCX Document Loader using python-docx.
Implements Adapter pattern for third-party library.
"""
from pathlib import Path
from docx import Document as DocxDocument

from ..core import Document, DocumentLoader, DocumentType


class DOCXLoader(DocumentLoader):
    """
    Adapter for python-docx library.
    Loads DOCX documents and extracts text content.
    """
    
    def __init__(self):
        self._supported_types = [DocumentType.DOCX]
    
    def load(self, path: str) -> Document:
        """
        Load DOCX and extract text.
        
        Args:
            path: Path to DOCX file.
            
        Returns:
            Document with extracted text.
            
        Raises:
            FileNotFoundError: If file doesn't exist.
            ValueError: If file is not a valid DOCX.
        """
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        if file_path.suffix.lower() != ".docx":
            raise ValueError(f"Expected .docx file, got: {file_path.suffix}")
        
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        content = "\n".join(text_parts)
        
        # Extract metadata
        metadata = {
            "filename": file_path.name,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        
        return Document(
            path=str(file_path.absolute()),
            type=DocumentType.DOCX,
            content=content,
            metadata=metadata
        )
    
    def supports(self, document_type: DocumentType) -> bool:
        """Check if loader supports DOCX type."""
        return document_type in self._supported_types
