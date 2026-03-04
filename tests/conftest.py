"""Pytest configuration and fixtures."""
import pytest
import os

# Fixed paths for test documents
_PDF_PATH = "/tmp/test_document_processor_sample.pdf"
_DOCX_PATH = "/tmp/test_document_processor_sample.docx"


def _create_test_files():
    """Create test documents if they don't exist."""
    # Create PDF
    if not os.path.exists(_PDF_PATH):
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), "Sample PDF content for testing.")
        doc.save(_PDF_PATH)
        doc.close()
    
    # Create DOCX
    if not os.path.exists(_DOCX_PATH):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Sample DOCX content for testing.")
        doc.add_paragraph("Second paragraph.")
        doc.save(_DOCX_PATH)


@pytest.fixture(scope='session', autouse=True)
def setup_test_documents():
    """Create test documents once per session."""
    _create_test_files()


@pytest.fixture
def sample_text():
    """Sample text for testing."""
    return """
    This is a sample document for testing the document processor.
    It contains multiple sentences and paragraphs.
    
    This is the second paragraph with more content.
    We need to test chunking and embedding functionality.
    
    The third paragraph concludes our test document.
    """


@pytest.fixture
def long_text():
    """Long text for testing chunking."""
    paragraphs = []
    for i in range(20):
        paragraphs.append(f"Paragraph {i}: " + " ".join([f"Sentence {j}." for j in range(5)]))
    return "\n\n".join(paragraphs)


@pytest.fixture(scope='session')
def session_sample_pdf():
    """Return path to sample PDF (session scope)."""
    return _PDF_PATH


@pytest.fixture(scope='session')
def session_sample_docx():
    """Return path to sample DOCX (session scope)."""
    return _DOCX_PATH
