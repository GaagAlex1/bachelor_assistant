"""Tests for document loaders."""
import pytest
from pathlib import Path

from document_processor.loaders import PDFLoader, DOCXLoader, LoaderFactory
from document_processor.core import DocumentType


class TestPDFLoader:
    """Tests for PDFLoader."""
    
    def test_load_pdf(self, tmp_path):
        """Test loading PDF file."""
        import fitz
        
        # Create sample PDF
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), "Test PDF content")
        doc.save(pdf_path)
        doc.close()
        
        loader = PDFLoader()
        document = loader.load(str(pdf_path))
        
        assert document.type == DocumentType.PDF
        assert "Test PDF content" in document.content
        assert document.metadata["pages"] == 1
    
    def test_load_nonexistent(self):
        """Test loading non-existent file."""
        loader = PDFLoader()
        
        with pytest.raises(FileNotFoundError):
            loader.load("nonexistent.pdf")
    
    def test_load_wrong_extension(self, tmp_path):
        """Test loading file with wrong extension."""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("Content")
        
        loader = PDFLoader()
        
        with pytest.raises(ValueError):
            loader.load(str(txt_path))
    
    def test_supports_method(self):
        """Test supports method."""
        loader = PDFLoader()
        
        assert loader.supports(DocumentType.PDF) is True
        assert loader.supports(DocumentType.DOCX) is False


class TestDOCXLoader:
    """Tests for DOCXLoader."""
    
    def test_load_docx(self, tmp_path):
        """Test loading DOCX file."""
        from docx import Document
        
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test DOCX content")
        doc.add_paragraph("Second paragraph")
        doc.save(docx_path)
        
        loader = DOCXLoader()
        document = loader.load(str(docx_path))
        
        assert document.type == DocumentType.DOCX
        assert "Test DOCX content" in document.content
        assert "Second paragraph" in document.content
        assert document.metadata["paragraphs"] == 2
    
    def test_load_docx_with_table(self, tmp_path):
        """Test loading DOCX with table."""
        from docx import Document
        
        docx_path = tmp_path / "test_table.docx"
        doc = Document()
        doc.add_paragraph("Document with table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        doc.save(docx_path)
        
        loader = DOCXLoader()
        document = loader.load(str(docx_path))
        
        assert document.type == DocumentType.DOCX
        assert document.metadata["tables"] == 1
        assert "Cell 1" in document.content
    
    def test_load_nonexistent(self):
        """Test loading non-existent file."""
        loader = DOCXLoader()
        
        with pytest.raises(FileNotFoundError):
            loader.load("nonexistent.docx")
    
    def test_supports_method(self):
        """Test supports method."""
        loader = DOCXLoader()
        
        assert loader.supports(DocumentType.DOCX) is True
        assert loader.supports(DocumentType.PDF) is False


class TestLoaderFactory:
    """Tests for LoaderFactory."""
    
    def test_get_loader_pdf(self):
        """Test getting PDF loader."""
        loader = LoaderFactory.get_loader(DocumentType.PDF)
        
        assert isinstance(loader, PDFLoader)
    
    def test_get_loader_docx(self):
        """Test getting DOCX loader."""
        loader = LoaderFactory.get_loader(DocumentType.DOCX)
        
        assert isinstance(loader, DOCXLoader)
    
    def test_get_loader_unsupported(self):
        """Test getting loader for unsupported type."""
        with pytest.raises(ValueError):
            LoaderFactory.get_loader("txt")
    
    def test_from_path_pdf(self):
        """Test getting loader from PDF path."""
        loader = LoaderFactory.from_path("document.pdf")
        
        assert isinstance(loader, PDFLoader)
    
    def test_from_path_docx(self):
        """Test getting loader from DOCX path."""
        loader = LoaderFactory.from_path("document.docx")
        
        assert isinstance(loader, DOCXLoader)
    
    def test_from_path_unsupported(self):
        """Test getting loader from unsupported path."""
        with pytest.raises(ValueError):
            LoaderFactory.from_path("document.txt")
    
    def test_supported_types(self):
        """Test getting supported types."""
        types = LoaderFactory.supported_types()
        
        assert DocumentType.PDF in types
        assert DocumentType.DOCX in types
    
    def test_register_custom_loader(self):
        """Test registering custom loader."""
        from document_processor.core import DocumentLoader, Document
        
        class CustomLoader(DocumentLoader):
            def load(self, path: str) -> Document:
                pass
            
            def supports(self, document_type: DocumentType) -> bool:
                return True
        
        # Register
        custom_type = DocumentType.PDF  # Use existing for test
        LoaderFactory.register_loader(custom_type, CustomLoader)
        
        loader = LoaderFactory.get_loader(custom_type)
        assert isinstance(loader, CustomLoader)
